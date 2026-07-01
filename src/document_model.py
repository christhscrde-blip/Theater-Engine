from __future__ import annotations

import hashlib
import re
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document

from .classifier import classify_texts, parse_speaker_line
from .models import ClassifiedParagraph, ParagraphType, TextSpan, TextSpanType
from .verifier import visible_text_hash

INLINE_STAGE_PATTERN = re.compile(r"\([^)]{1,160}\)")


@dataclass(frozen=True)
class DocumentParagraph:
    """A single paragraph in the internal, text-preserving document model."""

    index: int
    text: str
    classification: ClassifiedParagraph
    spans: tuple[TextSpan, ...] = field(default_factory=tuple)

    @property
    def reconstructed_text(self) -> str:
        return "".join(span.text for span in self.spans)

    @property
    def has_integrity(self) -> bool:
        return self.reconstructed_text == self.text

    @property
    def needs_manual_review(self) -> bool:
        return (
            self.classification.type == ParagraphType.UNCLASSIFIED
            or "needs_manual_review" in self.classification.flags
        )


@dataclass(frozen=True)
class DocumentModel:
    """Internal representation used between classification and formatting."""

    source_file: str
    visible_text_sha256: str
    paragraphs: tuple[DocumentParagraph, ...]

    @property
    def paragraph_count(self) -> int:
        return len(self.paragraphs)

    @property
    def visible_text(self) -> str:
        return "\n".join(paragraph.text for paragraph in self.paragraphs)

    @property
    def has_integrity(self) -> bool:
        expected_hash = hashlib.sha256(self.visible_text.encode("utf-8")).hexdigest()
        return expected_hash == self.visible_text_sha256 and all(
            paragraph.has_integrity for paragraph in self.paragraphs
        )

    @property
    def manual_review_paragraphs(self) -> tuple[DocumentParagraph, ...]:
        return tuple(
            paragraph for paragraph in self.paragraphs if paragraph.needs_manual_review
        )


def build_document_model_from_docx(docx_path: str | Path) -> DocumentModel:
    path = Path(docx_path)
    doc = Document(str(path))
    texts = [paragraph.text for paragraph in doc.paragraphs]
    return build_document_model_from_texts(
        texts,
        source_file=str(path),
        expected_visible_text_sha256=visible_text_hash(path),
    )


def build_document_model_from_texts(
    texts: list[str] | tuple[str, ...],
    source_file: str = "",
    expected_visible_text_sha256: str | None = None,
) -> DocumentModel:
    normalized_texts = [text if text is not None else "" for text in texts]
    classifications = classify_texts(normalized_texts)
    paragraphs = tuple(
        _build_document_paragraph(text, classification)
        for text, classification in zip(normalized_texts, classifications, strict=True)
    )
    visible_text = "\n".join(normalized_texts)
    text_hash = (
        expected_visible_text_sha256
        or hashlib.sha256(visible_text.encode("utf-8")).hexdigest()
    )
    model = DocumentModel(
        source_file=source_file,
        visible_text_sha256=text_hash,
        paragraphs=paragraphs,
    )
    _assert_model_integrity(model)
    return model


def _build_document_paragraph(
    text: str, classification: ClassifiedParagraph
) -> DocumentParagraph:
    spans = _build_spans(text, classification)
    paragraph = DocumentParagraph(
        index=classification.index,
        text=text,
        classification=classification,
        spans=spans,
    )
    if not paragraph.has_integrity:
        raise ValueError(
            f"DocumentModel integrity error in paragraph {classification.index}: "
            "span text does not reconstruct paragraph text"
        )
    return paragraph


def _build_spans(
    text: str, classification: ClassifiedParagraph
) -> tuple[TextSpan, ...]:
    if text == "":
        return (TextSpan(type=TextSpanType.PLAIN, text=""),)

    if classification.type == ParagraphType.SPEAKER_WITH_REPLIQUE:
        return _split_speaker_prefix_spans(text, classification, split_replique=True)
    if classification.type == ParagraphType.SPEAKER_WITH_STAGE:
        return _split_inline_stage_spans(
            text, TextSpanType.SPEAKER, speaker=classification.speaker
        )

    base_type = _base_span_type(classification.type)
    if classification.type == ParagraphType.REPLIQUE:
        return _split_inline_stage_spans(
            text, base_type, speaker=classification.speaker
        )

    return (TextSpan(type=base_type, text=text, speaker=classification.speaker),)


def _base_span_type(paragraph_type: ParagraphType) -> TextSpanType:
    if paragraph_type in {ParagraphType.SPEAKER, ParagraphType.SPEAKER_WITH_STAGE}:
        return TextSpanType.SPEAKER
    if paragraph_type in {ParagraphType.REPLIQUE, ParagraphType.SPEAKER_WITH_REPLIQUE}:
        return TextSpanType.REPLIQUE
    if paragraph_type == ParagraphType.STAGE_DIRECTION:
        return TextSpanType.STAGE_DIRECTION
    return TextSpanType.PLAIN


def _split_inline_stage_spans(
    text: str, fallback_type: TextSpanType, speaker: str = ""
) -> tuple[TextSpan, ...]:
    spans: list[TextSpan] = []
    cursor = 0
    for match in INLINE_STAGE_PATTERN.finditer(text):
        if match.start() > cursor:
            spans.append(
                TextSpan(
                    type=fallback_type,
                    text=text[cursor : match.start()],
                    speaker=speaker,
                )
            )
        spans.append(TextSpan(type=TextSpanType.INLINE_STAGE, text=match.group(0)))
        cursor = match.end()
    if cursor < len(text):
        spans.append(TextSpan(type=fallback_type, text=text[cursor:], speaker=speaker))
    return tuple(spans) or (TextSpan(type=fallback_type, text=text, speaker=speaker),)


def _split_speaker_prefix_spans(
    text: str, classification: ClassifiedParagraph, split_replique: bool
) -> tuple[TextSpan, ...]:
    parsed = parse_speaker_line(text)
    if not parsed:
        return _single_replique_span(text, classification)

    raw_name = parsed["raw"]
    name_start = text.find(raw_name)
    if name_start < 0:
        return _single_replique_span(text, classification)

    spans: list[TextSpan] = []
    if name_start:
        spans.append(TextSpan(type=TextSpanType.PLAIN, text=text[:name_start]))

    cursor = name_start + len(raw_name)
    stage_inline = parsed["stage_inline"]
    stage_match = (
        re.match(r"\s*" + re.escape(stage_inline), text[cursor:])
        if stage_inline
        else None
    )
    if stage_match:
        speaker_text = text[name_start:cursor]
        spans.append(
            TextSpan(
                type=TextSpanType.SPEAKER,
                text=speaker_text,
                speaker=classification.speaker,
            )
        )
        spans.append(
            TextSpan(type=TextSpanType.INLINE_STAGE, text=stage_match.group(0))
        )
        cursor += stage_match.end()
    else:
        separator_match = re.match(r"\s*[\.:]\s*", text[cursor:])
        if not separator_match:
            return _single_replique_span(text, classification)
        cursor += separator_match.end()
        speaker_text = text[name_start:cursor]
        spans.append(
            TextSpan(
                type=TextSpanType.SPEAKER,
                text=speaker_text,
                speaker=classification.speaker,
            )
        )

    if split_replique and cursor < len(text):
        spans.extend(
            _split_inline_stage_spans(
                text[cursor:], TextSpanType.REPLIQUE, speaker=classification.speaker
            )
        )
    elif cursor < len(text):
        spans.append(TextSpan(type=TextSpanType.PLAIN, text=text[cursor:]))

    if "".join(span.text for span in spans) != text:
        return _single_replique_span(text, classification)
    return tuple(spans)


def _single_replique_span(
    text: str, classification: ClassifiedParagraph
) -> tuple[TextSpan, ...]:
    return (
        TextSpan(type=TextSpanType.REPLIQUE, text=text, speaker=classification.speaker),
    )


def _assert_model_integrity(model: DocumentModel) -> None:
    if not model.has_integrity:
        raise ValueError("DocumentModel integrity error: visible text hash mismatch")
